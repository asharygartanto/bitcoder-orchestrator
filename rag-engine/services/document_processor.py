import os
import re
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from config import get_settings

settings = get_settings()


class DocumentProcessor:
    def __init__(self):
        self.upload_dir = settings.UPLOAD_DIR

    def extract_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._extract_docx(file_path)
        elif ext in (".txt", ".md"):
            return self._extract_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _extract_pdf(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n\n".join(texts)

    def _extract_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    texts.append(row_text)
        return "\n\n".join(texts)

    def _extract_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        try:
            import json
            data = json.loads(raw)
            if isinstance(data, dict) and "text" in data:
                return data["text"]
        except (json.JSONDecodeError, ValueError):
            pass
        return raw

    def process_document(self, file_path: str) -> list[dict]:
        text = self.extract_text(file_path)
        chunks = self.chunk_text(text)

        source_type = "document"
        source_url = ""
        try:
            import json
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw = f.read()
            data = json.loads(raw)
            if isinstance(data, dict):
                source_type = data.get("source_type", "document")
                source_url = data.get("source_url", "")
        except (json.JSONDecodeError, ValueError):
            pass

        for chunk in chunks:
            chunk["source_type"] = source_type
            chunk["source_url"] = source_url

        return chunks

    def process_document(self, file_path: str) -> list[dict]:
        text = self.extract_text(file_path)
        chunks = self.chunk_text(text)
        return chunks
